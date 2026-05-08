"""
CV Parser Service
Extracts plain text from PDF and DOCX files.
"""
import io
from pathlib import Path


def extract_text(file_path: str) -> str:
    """Extract text from PDF or DOCX file."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _extract_from_pdf(file_path: str) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()
        return "\n".join(text_parts).strip()
    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: pip install PyMuPDF")


def _extract_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
